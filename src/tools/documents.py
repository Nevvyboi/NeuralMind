"""
GroundZero AI - Document Understanding Engine
=============================================

State-of-the-art document processing:
1. Read ANY file type (PDF, Excel, CSV, Word, images, text, JSON, etc.)
2. Extract and understand content
3. Chunk large documents intelligently
4. Build document index for Q&A
5. Answer questions about document content
6. Multi-document analysis
7. Table/data extraction
8. Image OCR
"""

import os
import json
import re
import csv
import io
import base64
import hashlib
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple, Union
from dataclasses import dataclass, field, asdict
from collections import defaultdict

try:
    from ..utils import get_data_path, ensure_dir, load_json, save_json, logger, timestamp, generate_id
except ImportError:
    from utils import get_data_path, ensure_dir, load_json, save_json, logger, timestamp, generate_id


# ============================================================================
# DOCUMENT TYPES
# ============================================================================

@dataclass
class DocumentChunk:
    """A chunk of document content."""
    id: str = field(default_factory=lambda: generate_id("chunk_"))
    content: str = ""
    chunk_index: int = 0
    start_char: int = 0
    end_char: int = 0
    page_number: Optional[int] = None
    section: Optional[str] = None
    metadata: Dict = field(default_factory=dict)
    
    def to_dict(self) -> Dict:
        return asdict(self)


@dataclass
class ExtractedTable:
    """An extracted table from a document."""
    id: str = field(default_factory=lambda: generate_id("table_"))
    headers: List[str] = field(default_factory=list)
    rows: List[List[Any]] = field(default_factory=list)
    page_number: Optional[int] = None
    source: str = ""
    
    def to_dataframe(self):
        """Convert to pandas DataFrame if available."""
        try:
            import pandas as pd
            return pd.DataFrame(self.rows, columns=self.headers)
        except ImportError:
            return {"headers": self.headers, "rows": self.rows}
    
    def to_csv(self) -> str:
        """Convert to CSV string."""
        output = io.StringIO()
        writer = csv.writer(output)
        if self.headers:
            writer.writerow(self.headers)
        writer.writerows(self.rows)
        return output.getvalue()
    
    def to_dict(self) -> Dict:
        return asdict(self)


@dataclass 
class Document:
    """A processed document with extracted content."""
    id: str = field(default_factory=lambda: generate_id("doc_"))
    filename: str = ""
    filepath: str = ""
    file_type: str = ""
    file_size: int = 0
    
    # Content
    raw_content: str = ""
    chunks: List[DocumentChunk] = field(default_factory=list)
    tables: List[ExtractedTable] = field(default_factory=list)
    
    # Metadata
    title: Optional[str] = None
    author: Optional[str] = None
    created_date: Optional[str] = None
    page_count: int = 0
    word_count: int = 0
    language: str = "en"
    
    # Processing info
    processed_at: str = field(default_factory=timestamp)
    processing_time: float = 0.0
    extraction_method: str = ""
    
    # For Q&A
    summary: str = ""
    key_topics: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return {
            **asdict(self),
            "chunks": [c.to_dict() for c in self.chunks],
            "tables": [t.to_dict() for t in self.tables],
        }
    
    def get_full_text(self) -> str:
        """Get full document text."""
        if self.chunks:
            return "\n\n".join([c.content for c in self.chunks])
        return self.raw_content
    
    def search(self, query: str) -> List[DocumentChunk]:
        """Search for relevant chunks."""
        query_lower = query.lower()
        query_words = set(query_lower.split())
        
        scored_chunks = []
        for chunk in self.chunks:
            content_lower = chunk.content.lower()
            
            # Score based on word overlap
            chunk_words = set(content_lower.split())
            overlap = len(query_words & chunk_words)
            
            # Bonus for exact phrase
            if query_lower in content_lower:
                overlap += 5
            
            if overlap > 0:
                scored_chunks.append((overlap, chunk))
        
        # Sort by score
        scored_chunks.sort(key=lambda x: x[0], reverse=True)
        return [c for _, c in scored_chunks[:10]]


# ============================================================================
# FILE READERS
# ============================================================================

class PDFReader:
    """Read PDF files."""
    
    @staticmethod
    def read(filepath: str) -> Tuple[str, Dict]:
        """Read PDF and extract text."""
        metadata = {"pages": 0, "method": ""}
        
        # Try PyMuPDF (fitz) first - best quality
        try:
            import fitz
            doc = fitz.open(filepath)
            
            text_parts = []
            for page_num, page in enumerate(doc):
                text = page.get_text()
                text_parts.append(f"[Page {page_num + 1}]\n{text}")
            
            metadata["pages"] = len(doc)
            metadata["method"] = "pymupdf"
            metadata["title"] = doc.metadata.get("title", "")
            metadata["author"] = doc.metadata.get("author", "")
            
            doc.close()
            return "\n\n".join(text_parts), metadata
            
        except ImportError:
            pass
        
        # Try PyPDF2
        try:
            import PyPDF2
            
            text_parts = []
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                metadata["pages"] = len(reader.pages)
                
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
                
                if reader.metadata:
                    metadata["title"] = reader.metadata.get("/Title", "")
                    metadata["author"] = reader.metadata.get("/Author", "")
            
            metadata["method"] = "pypdf2"
            return "\n\n".join(text_parts), metadata
            
        except ImportError:
            pass
        
        # Try pdfplumber
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(filepath) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
            
            metadata["method"] = "pdfplumber"
            return "\n\n".join(text_parts), metadata
            
        except ImportError:
            pass
        
        raise ImportError("No PDF library available. Install: pip install pymupdf PyPDF2 pdfplumber")


class ExcelReader:
    """Read Excel files."""
    
    @staticmethod
    def read(filepath: str) -> Tuple[str, Dict, List[ExtractedTable]]:
        """Read Excel file and extract content."""
        metadata = {"sheets": [], "method": ""}
        tables = []
        text_parts = []
        
        # Try openpyxl
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(filepath, data_only=True)
            metadata["sheets"] = wb.sheetnames
            metadata["method"] = "openpyxl"
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"[Sheet: {sheet_name}]")
                
                # Extract as table
                rows = []
                headers = []
                
                for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    
                    if row_idx == 0:
                        headers = row_data
                    else:
                        rows.append(row_data)
                    
                    text_parts.append(" | ".join(row_data))
                
                if headers or rows:
                    tables.append(ExtractedTable(
                        headers=headers,
                        rows=rows,
                        source=f"{filepath}:{sheet_name}",
                    ))
                
                text_parts.append("")
            
            wb.close()
            return "\n".join(text_parts), metadata, tables
            
        except ImportError:
            pass
        
        # Try pandas
        try:
            import pandas as pd
            
            xlsx = pd.ExcelFile(filepath)
            metadata["sheets"] = xlsx.sheet_names
            metadata["method"] = "pandas"
            
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                text_parts.append(f"[Sheet: {sheet_name}]")
                text_parts.append(df.to_string())
                text_parts.append("")
                
                tables.append(ExtractedTable(
                    headers=list(df.columns),
                    rows=df.values.tolist(),
                    source=f"{filepath}:{sheet_name}",
                ))
            
            return "\n".join(text_parts), metadata, tables
            
        except ImportError:
            pass
        
        raise ImportError("No Excel library available. Install: pip install openpyxl pandas")


class CSVReader:
    """Read CSV/TSV files."""
    
    @staticmethod
    def read(filepath: str, delimiter: str = None) -> Tuple[str, Dict, List[ExtractedTable]]:
        """Read CSV file."""
        metadata = {"rows": 0, "columns": 0}
        
        # Auto-detect delimiter
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            sample = f.read(4096)
            if delimiter is None:
                if '\t' in sample:
                    delimiter = '\t'
                else:
                    delimiter = ','
        
        # Try pandas first
        try:
            import pandas as pd
            
            df = pd.read_csv(filepath, delimiter=delimiter)
            metadata["rows"] = len(df)
            metadata["columns"] = len(df.columns)
            metadata["column_names"] = list(df.columns)
            
            table = ExtractedTable(
                headers=list(df.columns),
                rows=df.values.tolist(),
                source=filepath,
            )
            
            return df.to_string(), metadata, [table]
            
        except ImportError:
            pass
        
        # Fallback to csv module
        rows = []
        headers = []
        
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f, delimiter=delimiter)
            for i, row in enumerate(reader):
                if i == 0:
                    headers = row
                else:
                    rows.append(row)
        
        metadata["rows"] = len(rows)
        metadata["columns"] = len(headers)
        metadata["column_names"] = headers
        
        table = ExtractedTable(headers=headers, rows=rows, source=filepath)
        
        text_parts = [" | ".join(headers)]
        text_parts.extend([" | ".join(row) for row in rows[:100]])  # Limit preview
        if len(rows) > 100:
            text_parts.append(f"... and {len(rows) - 100} more rows")
        
        return "\n".join(text_parts), metadata, [table]


class WordReader:
    """Read Word documents."""
    
    @staticmethod
    def read(filepath: str) -> Tuple[str, Dict]:
        """Read Word document."""
        metadata = {"paragraphs": 0, "method": ""}
        
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(filepath)
            metadata["method"] = "python-docx"
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            metadata["paragraphs"] = len(text_parts)
            
            # Extract tables
            for table in doc.tables:
                text_parts.append("\n[Table]")
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    text_parts.append(" | ".join(cells))
            
            # Try to get metadata
            try:
                metadata["title"] = doc.core_properties.title
                metadata["author"] = doc.core_properties.author
            except:
                pass
            
            return "\n\n".join(text_parts), metadata
            
        except ImportError:
            raise ImportError("python-docx not available. Install: pip install python-docx")


class ImageReader:
    """Read and OCR images."""
    
    @staticmethod
    def read(filepath: str) -> Tuple[str, Dict]:
        """Read image and extract text via OCR."""
        metadata = {"method": "", "dimensions": None}
        
        # Try pytesseract
        try:
            from PIL import Image
            import pytesseract
            
            img = Image.open(filepath)
            metadata["dimensions"] = img.size
            metadata["method"] = "pytesseract"
            
            text = pytesseract.image_to_string(img)
            return text, metadata
            
        except ImportError:
            pass
        
        # Try easyocr
        try:
            import easyocr
            
            reader = easyocr.Reader(['en'])
            results = reader.readtext(filepath)
            
            metadata["method"] = "easyocr"
            text = "\n".join([r[1] for r in results])
            return text, metadata
            
        except ImportError:
            pass
        
        # Return placeholder if no OCR available
        metadata["method"] = "none"
        return f"[Image file: {filepath} - OCR not available. Install: pip install pytesseract Pillow]", metadata


class JSONReader:
    """Read JSON files."""
    
    @staticmethod
    def read(filepath: str) -> Tuple[str, Dict]:
        """Read JSON file."""
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        metadata = {
            "type": type(data).__name__,
            "keys": list(data.keys()) if isinstance(data, dict) else None,
            "length": len(data) if isinstance(data, (list, dict)) else None,
        }
        
        # Pretty format
        text = json.dumps(data, indent=2, default=str)
        return text, metadata


class TextReader:
    """Read plain text files."""
    
    @staticmethod
    def read(filepath: str) -> Tuple[str, Dict]:
        """Read text file."""
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        lines = content.split('\n')
        metadata = {
            "lines": len(lines),
            "characters": len(content),
            "words": len(content.split()),
        }
        
        return content, metadata


# ============================================================================
# DOCUMENT PROCESSOR
# ============================================================================

class DocumentProcessor:
    """
    Process and understand documents of any type.
    """
    
    # Supported file types
    SUPPORTED_TYPES = {
        # Documents
        '.pdf': 'pdf',
        '.docx': 'word',
        '.doc': 'word',
        
        # Spreadsheets
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.csv': 'csv',
        '.tsv': 'csv',
        
        # Text
        '.txt': 'text',
        '.md': 'text',
        '.py': 'text',
        '.js': 'text',
        '.html': 'text',
        '.xml': 'text',
        '.json': 'json',
        '.yaml': 'text',
        '.yml': 'text',
        
        # Images
        '.png': 'image',
        '.jpg': 'image',
        '.jpeg': 'image',
        '.gif': 'image',
        '.bmp': 'image',
        '.tiff': 'image',
    }
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Document storage
        self.documents: Dict[str, Document] = {}
        self.storage_path = get_data_path("documents")
        ensure_dir(self.storage_path)
        
        self._load_index()
    
    def _load_index(self):
        """Load document index."""
        index_path = self.storage_path / "index.json"
        if index_path.exists():
            index = load_json(index_path)
            # Just load metadata, not full content
            for doc_id, doc_data in index.get("documents", {}).items():
                self.documents[doc_id] = doc_data  # Store as dict for now
    
    def _save_index(self):
        """Save document index."""
        index = {
            "documents": {
                doc_id: {
                    "id": doc.id if isinstance(doc, Document) else doc.get("id"),
                    "filename": doc.filename if isinstance(doc, Document) else doc.get("filename"),
                    "filepath": doc.filepath if isinstance(doc, Document) else doc.get("filepath"),
                    "file_type": doc.file_type if isinstance(doc, Document) else doc.get("file_type"),
                    "processed_at": doc.processed_at if isinstance(doc, Document) else doc.get("processed_at"),
                }
                for doc_id, doc in self.documents.items()
            },
            "updated_at": timestamp(),
        }
        save_json(self.storage_path / "index.json", index)
    
    def process(self, filepath: str) -> Document:
        """
        Process a document and extract content.
        
        Args:
            filepath: Path to the document
        
        Returns:
            Processed Document object
        """
        start_time = datetime.now()
        path = Path(filepath)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {filepath}")
        
        # Get file type
        ext = path.suffix.lower()
        file_type = self.SUPPORTED_TYPES.get(ext, 'unknown')
        
        if file_type == 'unknown':
            # Try to read as text
            file_type = 'text'
        
        # Create document
        doc = Document(
            filename=path.name,
            filepath=str(path.absolute()),
            file_type=file_type,
            file_size=path.stat().st_size,
        )
        
        # Read content based on type
        tables = []
        metadata = {}
        
        try:
            if file_type == 'pdf':
                content, metadata = PDFReader.read(filepath)
            elif file_type == 'excel':
                content, metadata, tables = ExcelReader.read(filepath)
            elif file_type == 'csv':
                content, metadata, tables = CSVReader.read(filepath)
            elif file_type == 'word':
                content, metadata = WordReader.read(filepath)
            elif file_type == 'image':
                content, metadata = ImageReader.read(filepath)
            elif file_type == 'json':
                content, metadata = JSONReader.read(filepath)
            else:
                content, metadata = TextReader.read(filepath)
            
            doc.raw_content = content
            doc.tables = tables
            doc.extraction_method = metadata.get("method", file_type)
            doc.page_count = metadata.get("pages", 0)
            doc.title = metadata.get("title")
            doc.author = metadata.get("author")
            
        except Exception as e:
            logger.error(f"Error reading {filepath}: {e}")
            doc.raw_content = f"Error reading file: {e}"
        
        # Calculate word count
        doc.word_count = len(doc.raw_content.split())
        
        # Chunk the document
        doc.chunks = self._chunk_document(doc.raw_content)
        
        # Extract key topics (simple keyword extraction)
        doc.key_topics = self._extract_topics(doc.raw_content)
        
        # Generate summary (first chunk + key info)
        doc.summary = self._generate_summary(doc)
        
        # Processing time
        doc.processing_time = (datetime.now() - start_time).total_seconds()
        
        # Store
        self.documents[doc.id] = doc
        
        # Save full document
        doc_path = self.storage_path / f"{doc.id}.json"
        save_json(doc_path, doc.to_dict())
        
        # Update index
        self._save_index()
        
        logger.info(f"Processed document: {doc.filename} ({doc.word_count} words, {len(doc.chunks)} chunks)")
        
        return doc
    
    def _chunk_document(self, text: str) -> List[DocumentChunk]:
        """Split document into chunks."""
        chunks = []
        
        # Split by paragraphs first
        paragraphs = re.split(r'\n\s*\n', text)
        
        current_chunk = ""
        current_start = 0
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # If adding this paragraph exceeds chunk size
            if len(current_chunk) + len(para) > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append(DocumentChunk(
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    start_char=current_start,
                    end_char=current_start + len(current_chunk),
                ))
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else ""
                current_chunk = overlap_text + "\n\n" + para
                current_start = current_start + len(current_chunk) - len(overlap_text)
                chunk_index += 1
            else:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
        
        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=current_start + len(current_chunk),
            ))
        
        return chunks
    
    def _extract_topics(self, text: str, max_topics: int = 10) -> List[str]:
        """Extract key topics from text."""
        # Simple keyword extraction
        words = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
        
        # Count frequencies
        word_counts = defaultdict(int)
        for word in words:
            if len(word) > 3:
                word_counts[word] += 1
        
        # Get top words
        sorted_words = sorted(word_counts.items(), key=lambda x: x[1], reverse=True)
        return [word for word, _ in sorted_words[:max_topics]]
    
    def _generate_summary(self, doc: Document) -> str:
        """Generate a simple summary."""
        parts = [f"Document: {doc.filename}"]
        parts.append(f"Type: {doc.file_type}")
        parts.append(f"Size: {doc.word_count} words")
        
        if doc.page_count:
            parts.append(f"Pages: {doc.page_count}")
        
        if doc.tables:
            parts.append(f"Tables: {len(doc.tables)}")
        
        if doc.key_topics:
            parts.append(f"Key topics: {', '.join(doc.key_topics[:5])}")
        
        # First 500 chars of content
        if doc.raw_content:
            preview = doc.raw_content[:500].replace('\n', ' ')
            parts.append(f"\nPreview: {preview}...")
        
        return "\n".join(parts)
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get a document by ID."""
        if doc_id in self.documents:
            doc = self.documents[doc_id]
            if isinstance(doc, Document):
                return doc
            
            # Load full document from disk
            doc_path = self.storage_path / f"{doc_id}.json"
            if doc_path.exists():
                data = load_json(doc_path)
                doc = Document(**{k: v for k, v in data.items() 
                                 if k not in ['chunks', 'tables']})
                doc.chunks = [DocumentChunk(**c) for c in data.get('chunks', [])]
                doc.tables = [ExtractedTable(**t) for t in data.get('tables', [])]
                self.documents[doc_id] = doc
                return doc
        
        return None
    
    def search_documents(self, query: str, doc_ids: List[str] = None) -> List[Tuple[Document, List[DocumentChunk]]]:
        """
        Search across documents for relevant content.
        
        Args:
            query: Search query
            doc_ids: Optional list of doc IDs to search (searches all if None)
        
        Returns:
            List of (document, matching_chunks) tuples
        """
        results = []
        
        search_docs = doc_ids or list(self.documents.keys())
        
        for doc_id in search_docs:
            doc = self.get_document(doc_id)
            if doc:
                matching_chunks = doc.search(query)
                if matching_chunks:
                    results.append((doc, matching_chunks))
        
        return results
    
    def get_context_for_query(self, query: str, doc_ids: List[str] = None, max_chunks: int = 5) -> str:
        """
        Get relevant document context for a query.
        
        This is what you use for Q&A - it returns the most relevant
        chunks to include in the prompt.
        """
        results = self.search_documents(query, doc_ids)
        
        context_parts = []
        chunk_count = 0
        
        for doc, chunks in results:
            for chunk in chunks:
                if chunk_count >= max_chunks:
                    break
                
                context_parts.append(f"[From: {doc.filename}]\n{chunk.content}")
                chunk_count += 1
            
            if chunk_count >= max_chunks:
                break
        
        return "\n\n---\n\n".join(context_parts)
    
    def list_documents(self) -> List[Dict]:
        """List all processed documents."""
        docs = []
        for doc_id, doc in self.documents.items():
            if isinstance(doc, Document):
                docs.append({
                    "id": doc.id,
                    "filename": doc.filename,
                    "file_type": doc.file_type,
                    "word_count": doc.word_count,
                    "chunks": len(doc.chunks),
                    "processed_at": doc.processed_at,
                })
            else:
                docs.append(doc)
        return docs
    
    def delete_document(self, doc_id: str) -> bool:
        """Delete a document."""
        if doc_id in self.documents:
            del self.documents[doc_id]
            
            doc_path = self.storage_path / f"{doc_id}.json"
            if doc_path.exists():
                doc_path.unlink()
            
            self._save_index()
            return True
        return False


# ============================================================================
# DOCUMENT Q&A ENGINE
# ============================================================================

class DocumentQA:
    """
    Question-answering over documents.
    
    This is the main interface for asking questions about document content.
    """
    
    def __init__(self, processor: DocumentProcessor, model_generate=None):
        self.processor = processor
        self.model_generate = model_generate
        
        # Active documents for current session
        self.active_documents: List[str] = []
    
    def load_document(self, filepath: str) -> Document:
        """Load and process a document."""
        doc = self.processor.process(filepath)
        self.active_documents.append(doc.id)
        return doc
    
    def load_documents(self, filepaths: List[str]) -> List[Document]:
        """Load multiple documents."""
        docs = []
        for fp in filepaths:
            try:
                doc = self.load_document(fp)
                docs.append(doc)
            except Exception as e:
                logger.error(f"Error loading {fp}: {e}")
        return docs
    
    def ask(self, question: str, doc_ids: List[str] = None) -> Dict:
        """
        Ask a question about the documents.
        
        Args:
            question: The question to answer
            doc_ids: Specific documents to search (uses active docs if None)
        
        Returns:
            Dict with answer, sources, and confidence
        """
        # Use active documents if not specified
        search_ids = doc_ids or self.active_documents
        
        if not search_ids:
            return {
                "answer": "No documents loaded. Please load documents first.",
                "sources": [],
                "confidence": 0.0,
            }
        
        # Get relevant context
        context = self.processor.get_context_for_query(question, search_ids, max_chunks=5)
        
        if not context:
            return {
                "answer": "I couldn't find relevant information in the documents.",
                "sources": [],
                "confidence": 0.0,
            }
        
        # Generate answer
        if self.model_generate:
            prompt = f"""Based on the following document excerpts, answer the question.

DOCUMENT CONTENT:
{context}

QUESTION: {question}

Answer based only on the information provided. If the answer isn't in the documents, say so."""
            
            answer = self.model_generate(prompt)
        else:
            # Without model, return context directly
            answer = f"Relevant information from documents:\n\n{context}"
        
        # Extract source documents
        sources = []
        for doc_id in search_ids:
            doc = self.processor.get_document(doc_id)
            if doc:
                sources.append({
                    "id": doc.id,
                    "filename": doc.filename,
                    "file_type": doc.file_type,
                })
        
        return {
            "answer": answer,
            "sources": sources,
            "context": context,
            "confidence": 0.8 if context else 0.0,
        }
    
    def summarize(self, doc_id: str = None) -> str:
        """Summarize a document or all active documents."""
        if doc_id:
            doc = self.processor.get_document(doc_id)
            if not doc:
                return "Document not found."
            docs = [doc]
        else:
            docs = [self.processor.get_document(did) for did in self.active_documents]
            docs = [d for d in docs if d]
        
        if not docs:
            return "No documents to summarize."
        
        summaries = []
        for doc in docs:
            summaries.append(f"## {doc.filename}\n{doc.summary}")
        
        return "\n\n".join(summaries)
    
    def get_tables(self, doc_id: str = None) -> List[ExtractedTable]:
        """Get all tables from documents."""
        tables = []
        
        doc_ids = [doc_id] if doc_id else self.active_documents
        
        for did in doc_ids:
            doc = self.processor.get_document(did)
            if doc:
                tables.extend(doc.tables)
        
        return tables
    
    def clear(self):
        """Clear active documents."""
        self.active_documents = []


# Export
__all__ = [
    'DocumentChunk', 'ExtractedTable', 'Document',
    'PDFReader', 'ExcelReader', 'CSVReader', 'WordReader', 'ImageReader',
    'DocumentProcessor', 'DocumentQA',
]
